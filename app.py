import streamlit as st
import json
import re
import io
from pathlib import Path
from datetime import date

BASE = Path(__file__).parent
RUBRIC = json.loads((BASE / "specs/rubric.json").read_text())
QUESTION_BANK = json.loads((BASE / "scoring/question_bank.json").read_text())
SIGNAL_OPTIONS = json.loads((BASE / "scoring/signal_options.json").read_text())

# Overrides for signals where the interview question doesn't describe what's being rated
FORM_HINTS = {
    "s1_1": "Rate how specifically your target AI use cases are defined, scoped, and prioritized",
    "s1_2": "Rate the formality and authority of executive sponsorship for this AI initiative",
    "s1_3": "Rate how well your stated AI ambition matches your organization's actual capacity, budget, and capabilities",
    "s1_4": "Rate whether measurable success metrics are defined for your AI initiative",
    "s1_5": "How defined is the implementation timeline and urgency drivers for this AI initiative?",
    "s2_2": "Rate the quality of your data for the domains required by your AI use cases",
    "s2_3": "Rate your interoperability maturity — FHIR/HL7 standards, API capabilities, and cross-system data exchange",
    "s3_3": "Rate your network infrastructure's readiness for AI workloads — bandwidth, latency for real-time inference, and security segmentation",
    "s3_5": "Can your infrastructure scale with AI workload growth, and do your disaster recovery and business continuity plans explicitly cover AI-dependent clinical and operational workflows?",
    "s4_3": "Rate the maturity of your AI-specific incident response readiness and cybersecurity posture",
    "s4_4": "Rate how thoroughly your regulatory compliance requirements have been assessed for AI workloads — including cloud-based AI processing, third-party model training restrictions, and FDA AI/ML guidance (this is an assessment of readiness, not a compliance certification)",
    "s4_5": "Rate your requirements around how AI vendors may USE, store, and share your data — training restrictions, retention limits, residency, and subprocessor controls",
    "s4_6": "Rate your vendor contract requirements around security controls, SLAs, penetration testing rights, and breach notification — distinct from data-use restrictions",
    "s7_2": "Rate how well-defined your AI infrastructure requirements are before vendor conversations begin",
    "s7_3": "Rate your organization's competency to evaluate AI vendors — including formal scoring criteria, reference checks, and ability to design and run vendor proof-of-concepts",
    "s8_1": "Rate whether a realistic budget is allocated, approved, or fundable — covering licensing, implementation, integration, training, and ongoing operations",
    "s8_2": "Rate how well your available budget matches your stated AI ambition",
    "s8_3": "Rate how comprehensively your budget model accounts for total cost of ownership — licensing, implementation, integration, training, ongoing operations, support, and technology refresh",
}

NA_OPTION = "Not applicable — this area has not been assessed or does not apply to our organization"

STATED_LEVEL_SCORES = {
    "none": 1.0, "vague": 1.5, "partial": 2.5, "defined": 3.5, "specific_with_metrics": 4.5,
}
LEVEL_KEYS = ["none", "vague", "partial", "defined", "specific_with_metrics"]
DOMAIN_IDS = [
    "strategic_intent_use_case_scope",
    "data_readiness_governance",
    "infrastructure_architecture_readiness",
    "security_privacy_compliance",
    "ai_governance_operating_model",
    "workflow_adoption_change_readiness",
    "procurement_vendor_readiness",
    "budget_investment_readiness",
]
CRITICAL_GATING = {"security_privacy_compliance", "data_readiness_governance", "ai_governance_operating_model"}
SECONDARY_GATING = {"infrastructure_architecture_readiness", "procurement_vendor_readiness", "budget_investment_readiness"}
PHI_ATTESTATION_TEXT = QUESTION_BANK["metadata"]["safety_attestation_text"]

# ── Reference design library ───────────────────────────────────────────────────
# Derived from reference_library/*.md — action-oriented gap-closure note per domain per vendor.
# Intent: if a domain scores Red/Yellow, show what the matched vendor offers to close the gap.

REF_VENDOR_NAMES = {
    "epic":   "Epic AI Ecosystem",
    "aws":    "AWS Health AI",
    "azure":  "Azure Cloud for Healthcare",
    "google": "Google Cloud Health AI",
}

REF_VENDOR_FILES = {
    "epic":   "epic_ai_ecosystem.md",
    "aws":    "aws_health_ai.md",
    "azure":  "azure_health_foundation.md",
    "google": "google_cloud_health_ai.md",
}

DOMAIN_OWNERS = {
    "strategic_intent_use_case_scope":      "CIO",
    "data_readiness_governance":            "CDO / Data Architect",
    "infrastructure_architecture_readiness":"VP IT Infrastructure",
    "security_privacy_compliance":          "CISO",
    "ai_governance_operating_model":        "CMIO / AI Governance Lead",
    "workflow_adoption_change_readiness":   "CMIO / Change Management",
    "procurement_vendor_readiness":         "Supply Chain / Legal",
    "budget_investment_readiness":          "CFO / CIO",
}

REF_DOMAIN_CAPABILITIES = {
    "epic": {
        "strategic_intent_use_case_scope":
            "Consolidate AI tool procurement under Epic App Orchard to reduce integration risk and leverage existing Epic governance.",
        "data_readiness_governance":
            "Activate Epic 2024 NLP module and validate FHIR R4 API enablement before any third-party AI data integration.",
        "infrastructure_architecture_readiness":
            "Resolve Epic operational stability issues before AI module activation — Epic AI requires stable EHR operations as a prerequisite.",
        "security_privacy_compliance":
            "Require individual HIPAA BAAs per App Orchard partner; review each partner's data flow documentation before activation.",
        "ai_governance_operating_model":
            "Establish AI governance committee authority to approve/block App Orchard activations before any new tool goes live.",
        "workflow_adoption_change_readiness":
            "Deploy DAX Copilot (ambient documentation) or Epic Cheers (patient engagement AI) as first AI workflow activation — lower integration risk within Epic.",
        "procurement_vendor_readiness":
            "Verify App Orchard certification status per vendor; require clinical validation evidence and equity/bias testing in procurement requirements.",
        "budget_investment_readiness":
            "Negotiate App Orchard and DAX licensing in next Epic contract renewal; DAX costs may be includable in existing Microsoft EA.",
    },
    "aws": {
        "strategic_intent_use_case_scope":
            "Define specific data workloads before AWS engagement — HealthLake and Bedrock deliver value when use case requirements are concrete, not exploratory.",
        "data_readiness_governance":
            "Execute AWS BAA and configure HealthLake as FHIR-native data layer; Comprehend Medical extracts structured clinical entities from unstructured text.",
        "infrastructure_architecture_readiness":
            "Deploy SageMaker health reference architecture for target use case (readmission, deterioration, revenue cycle) — requires SI partner if internal data engineering is thin.",
        "security_privacy_compliance":
            "Confirm HIPAA eligibility of every AWS service in scope; create a service-specific HIPAA configuration checklist before any PHI workloads begin.",
        "ai_governance_operating_model":
            "Assign internal AI governance owner before AWS AI deployment — AWS infrastructure does not include clinical model validation or post-deployment monitoring.",
        "workflow_adoption_change_readiness":
            "Bedrock + Transcribe Medical provides ambient documentation infrastructure but requires custom workflow integration; plan 3–6 months of integration effort.",
        "procurement_vendor_readiness":
            "Include cost management controls (budget alerts, cost anomaly detection) and SI engagement plan in AWS procurement requirements.",
        "budget_investment_readiness":
            "AWS consumption pricing requires reserved instance planning for steady-state workloads; set cost anomaly detection thresholds from day one.",
    },
    "azure": {
        "strategic_intent_use_case_scope":
            "Leverage Microsoft EA to bundle Azure health AI services; DAX Copilot is the primary near-term AI activation path for Microsoft-aligned health systems.",
        "data_readiness_governance":
            "Azure Health Data Services (AHDS) provides FHIR-native aggregation — priority activation if multi-EHR data harmonization is blocking AI workloads.",
        "infrastructure_architecture_readiness":
            "Validate Azure tenant configuration and data residency settings before DAX Copilot activation; leverage existing Azure infrastructure investments.",
        "security_privacy_compliance":
            "Confirm Microsoft BAA covers all health AI services in scope; request subprocessor list and DAX data residency commitment before go-live.",
        "ai_governance_operating_model":
            "Create a physician AI use policy before DAX rollout; Microsoft adoption analytics support monitoring but governance must be owned by the health system.",
        "workflow_adoption_change_readiness":
            "DAX Copilot has the largest health system customer base for ambient documentation; physician adoption rate and change management are the primary implementation risks.",
        "procurement_vendor_readiness":
            "DAX procurement requires Nuance/Microsoft licensing, Azure tenant setup, and change management plan — build all three into RFP requirements.",
        "budget_investment_readiness":
            "DAX per-provider licensing scales with physician count; explore EA bundling with Microsoft before standalone DAX contract.",
    },
    "google": {
        "strategic_intent_use_case_scope":
            "Google Cloud is best matched to data-mature orgs with defined population health or research AI workloads; resolve data readiness gaps before Google engagement.",
        "data_readiness_governance":
            "Healthcare Data Engine (HDE) provides FHIR-native de-identification and multi-EHR harmonization — priority for orgs with complex data environments blocking AI workloads.",
        "infrastructure_architecture_readiness":
            "Vertex AI + MedLM requires internal data science capacity or SI partner; plan significant data engineering investment before model development begins.",
        "security_privacy_compliance":
            "Verify HIPAA eligibility per Google service against current Google HIPAA-eligible products list; strong de-identification reduces PHI risk in analytics workflows.",
        "ai_governance_operating_model":
            "Assign data science governance owner and model validation process before any Vertex AI or MedLM deployment.",
        "workflow_adoption_change_readiness":
            "Care Studio enables clinical data surfacing in complex multi-system environments; ambient documentation is less mature than DAX — confirm use case fit before procurement.",
        "procurement_vendor_readiness":
            "Google Cloud RFP requires per-service BAA coverage, data residency commitment, and SI partner engagement plan.",
        "budget_investment_readiness":
            "Google Cloud data pipeline and Vertex AI costs require proactive cost management planning; most cost-effective when a data platform investment is already committed.",
    },
}


def _truncate(text, max_chars):
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rsplit(" ", 1)[0] + "…"


def _build_domain_questions():
    qb_domains = QUESTION_BANK["domains"]
    result = {}
    for i, domain_id in enumerate(DOMAIN_IDS):
        dom = qb_domains.get(f"domain_{i+1}", {})
        qs = dom.get("questions", {})
        result[domain_id] = {
            "name": dom.get("name", domain_id),
            "questions": [
                {
                    "signal_id": q["signal_id"],
                    "signal_name": q["signal_name"],
                    "hint": FORM_HINTS.get(q["signal_id"]) or _truncate(q["question"].split("?")[0], 160),
                    "evidence_types": q.get("evidenceValidation", {}).get("acceptedEvidenceTypes", []),
                }
                for q in qs.values()
            ],
        }
    return result


# ── RFP domain requirements ────────────────────────────────────────────────────
# Specific procurement language for IT and procurement teams, keyed by domain.
# Used in the Word doc RFP brief — shown only for Red/Yellow domains.

RFP_DOMAIN_REQUIREMENTS = {
    "strategic_intent_use_case_scope": {
        "heading": "Use Case Fit & Deployment Experience",
        "items": [
            "Vendor must demonstrate prior production deployment experience with the specific AI use case(s) in scope at comparable health system organizations (similar size, EHR platform, and care setting).",
            "Vendor must provide a minimum of two (2) health system reference customers currently using the proposed solution in production — references must be available for a reference call.",
            "Vendor must articulate in writing how the proposed solution addresses each AI use case listed in this RFP — responses must be use-case-specific, not generic.",
            "Vendor must provide a phased implementation plan with defined go-live milestones and advancement criteria between phases.",
            "Vendor must disclose any known limitations or contraindications of the proposed solution for the stated use cases.",
        ],
    },
    "data_readiness_governance": {
        "heading": "Data Requirements & Governance",
        "items": [
            "Vendor must provide complete data flow documentation: what data is ingested, where it is processed, where outputs are stored, and who (including all subprocessors) has access at each stage.",
            "Vendor must identify all subprocessors who receive, process, or store PHI on behalf of the health system — this list must be included in the BAA or as an executed BAA exhibit.",
            "Vendor must support FHIR R4 data exchange or provide a certified, documented integration path with the health system's EHR platform and active version.",
            "Vendor must provide data minimization documentation: what PHI is required for the solution to function vs. what is optional.",
            "Vendor must provide data retention, deletion, and de-identification policies — deletion timelines must be enforceable and auditable.",
            "Vendor must describe data quality requirements: minimum data quality standards the health system's input data must meet for the AI to function as specified.",
        ],
    },
    "infrastructure_architecture_readiness": {
        "heading": "Infrastructure & Architecture Requirements",
        "items": [
            "Vendor must provide a reference architecture diagram showing all integration touchpoints with the health system's EHR, network, and cloud environment.",
            "Vendor must specify compute requirements in detail: on-premises vs. cloud, minimum hardware or cloud instance specifications, cloud region requirements, and any GPU or specialized compute dependencies.",
            "Vendor must document network requirements: minimum bandwidth, latency thresholds for real-time inference, and all port and firewall configuration requirements.",
            "Vendor must describe disaster recovery and business continuity provisions specifically for AI-dependent clinical and operational workflows.",
            "Vendor must provide scalability documentation: how the solution handles workload growth, what triggers capacity scaling, and what the health system must provision to enable scaling.",
            "Vendor must describe integration approach for the health system's specific EHR platform and active release version.",
        ],
    },
    "security_privacy_compliance": {
        "heading": "Security, Privacy & Compliance Requirements",
        "items": [
            "Vendor must execute a HIPAA Business Associate Agreement (BAA) before any PHI is shared — BAA must be signed prior to the start of any proof-of-concept involving real patient data.",
            "Vendor must provide a current SOC 2 Type II report or equivalent third-party security certification dated within 18 months of contract execution.",
            "Vendor must provide most recent penetration testing results (within 12 months) and documented remediation for any high or critical findings.",
            "Vendor must document breach notification procedures: notification timeline (72-hour minimum), notification method, scope, and responsible parties.",
            "Vendor must provide a complete, current subprocessor list and commit to notifying the health system of any material subprocessor changes no less than 30 days in advance.",
            "Vendor must confirm data residency in writing: where PHI is stored and processed; US-only residency must be explicitly committed in the BAA if required by the health system.",
            "Vendor must describe encryption specifications: in-transit (TLS version minimum), at-rest (algorithm and key management), and any data-in-use protections.",
            "Vendor must provide a vulnerability disclosure policy and document the health system's rights to conduct security assessments.",
        ],
    },
    "ai_governance_operating_model": {
        "heading": "AI Governance & Model Transparency Requirements",
        "items": [
            "Vendor must provide a model card or equivalent documentation for each AI model proposed: training data sources and demographics, intended use, known limitations, and explicitly out-of-scope uses.",
            "Vendor must document the AI model update process: how and when models are updated or retrained, what customer notification is required before updates, and whether customers can opt out of or delay model updates.",
            "Vendor must provide equity and bias testing methodology and results for the specific model(s) proposed — evaluation must include demographic subgroups relevant to the health system's patient population.",
            "Vendor must describe post-deployment monitoring capabilities: drift detection, usage analytics, adverse event reporting mechanism, and process for the health system to surface concerns.",
            "Vendor must disclose FDA regulatory status for all clinical AI tools proposed — 510(k) clearance number, De Novo authorization, or exemption basis must be documented.",
            "Vendor must provide the health system with contractual audit rights: right to audit AI model performance, data handling, and subprocessor compliance on reasonable notice.",
        ],
    },
    "workflow_adoption_change_readiness": {
        "heading": "Workflow Integration & Clinical Adoption Requirements",
        "items": [
            "Vendor must provide a detailed implementation and workflow integration plan specifying how the solution surfaces within existing EHR and clinical workflows — core functions must not require clinicians to access a separate application.",
            "Vendor must provide adoption metrics from at least two comparable health system deployments: user adoption rate at 30/60/90 days, time to stable utilization, and documented barriers to adoption and how they were resolved.",
            "Vendor must describe training resources: curriculum, delivery method (on-site, virtual, asynchronous), and estimated time commitment for clinical and administrative staff.",
            "Vendor must describe change management support resources available during and after go-live, including super-user program and ongoing support model.",
            "Vendor must provide a formal rollback plan: steps and timeline to safely disable the AI tool if clinical or operational issues arise post go-live.",
            "Vendor must document the escalation path for clinical or patient safety concerns arising from AI tool behavior.",
        ],
    },
    "procurement_vendor_readiness": {
        "heading": "Procurement Process & Vendor Qualification Requirements",
        "items": [
            "Vendor must provide a detailed, itemized pricing model: licensing structure, per-unit costs (per provider, per patient, per transaction), volume discount thresholds, and annual price escalation caps.",
            "Vendor must provide a sample Master Services Agreement (MSA) — the health system will conduct legal review and vendor must be willing to negotiate standard terms.",
            "Vendor must provide a formal proof-of-concept (POC) proposal: scope, timeline, success criteria, data requirements, and cost — success criteria must be agreed upon in writing before the POC begins.",
            "Vendor must provide minimum two reference contacts from comparable health system procurements — similar size, EHR platform, and use case scope.",
            "Vendor must disclose any pending litigation, regulatory action, enforcement action, or material ownership or corporate structure changes that could affect the proposed solution.",
            "Vendor must disclose any planned product sunsetting, end-of-life timelines, or material roadmap changes within the proposed contract term.",
        ],
    },
    "budget_investment_readiness": {
        "heading": "Total Cost of Ownership & Financial Requirements",
        "items": [
            "Vendor must provide a total cost of ownership (TCO) model for years 1, 2, and 3: all licensing, implementation, integration, training, support, and technology refresh costs must be itemized separately.",
            "Vendor must itemize one-time implementation and professional services costs separately from annual recurring fees.",
            "Vendor must disclose all costs not included in base pricing: additional storage, API volume overages, out-of-scope professional services, new-staff training, and upgrade costs.",
            "Vendor must provide ROI data or outcome case studies from comparable health system deployments — methodology, baseline assumptions, and measurement timeframe must be documented.",
            "Vendor must describe contract termination provisions: data export rights, transition support availability, and any termination fees or penalties.",
            "Vendor must describe pricing at scale: how pricing changes if the health system expands to additional departments, facilities, or use cases during the contract term.",
        ],
    },
}


def _color(score):
    if score <= 2.4:
        return "Red"
    if score <= 3.4:
        return "Yellow"
    return "Green"


def _select_vendors(ehr, cloud):
    """Return 1-2 reference design vendor keys matched to EHR + cloud profile."""
    vendors = []
    if "Epic" in ehr:
        vendors.append("epic")
    if "Azure" in cloud or "Microsoft" in cloud:
        vendors.append("azure")
    elif "AWS" in cloud or "Amazon" in cloud:
        vendors.append("aws")
    elif "Google" in cloud:
        vendors.append("google")
    if not vendors:
        vendors.append("aws")
    return vendors[:2]


def _domain_headline(color, score, domain_id, n_weak, n_total):
    labels = {
        "strategic_intent_use_case_scope":       ("AI use case strategy",         "use case definition and executive mandate"),
        "data_readiness_governance":             ("Data foundation",               "data quality, governance, and interoperability"),
        "infrastructure_architecture_readiness": ("Infrastructure readiness",      "compute, cloud, and integration architecture"),
        "security_privacy_compliance":           ("Security and compliance posture","security controls and privacy governance"),
        "ai_governance_operating_model":         ("AI governance",                 "governance structure and operating model accountability"),
        "workflow_adoption_change_readiness":    ("Clinical adoption readiness",   "workflow integration and change management capacity"),
        "procurement_vendor_readiness":          ("Procurement readiness",         "vendor evaluation and RFP capability"),
        "budget_investment_readiness":           ("Budget readiness",              "investment alignment and total cost of ownership planning"),
    }
    area, gap_area = labels.get(domain_id, ("This domain", "this area"))
    if color == "Green":
        return f"{area} is well-developed — foundation is in place to proceed with AI procurement in this area."
    elif color == "Yellow":
        if n_weak <= 1:
            return f"{area} has strong foundations but targeted gaps must be closed before vendor engagement."
        return f"{area} has partial capability; {n_weak} of {n_total} signals require maturation before AI procurement can proceed."
    else:
        if n_weak == n_total:
            return f"{area} lacks foundational elements — {gap_area} gaps will block AI deployment until resolved."
        return f"{area} has significant gaps in {gap_area}; foundational work is required before vendor engagement or procurement."


def _decision_readiness_status(domain_scores):
    colors = {d: s["color"] for d, s in domain_scores.items()}
    critical_red = any(colors.get(d) == "Red" for d in CRITICAL_GATING)
    secondary_red = any(colors.get(d) == "Red" for d in SECONDARY_GATING)
    total_red = sum(1 for c in colors.values() if c == "Red")
    total_green = sum(1 for c in colors.values() if c == "Green")
    critical_evidence_ok = all(
        domain_scores.get(d, {}).get("evidence_confidence_score", 0) >= 3 for d in CRITICAL_GATING
    )
    if critical_red or total_red >= 2:
        key = "1_not_ready"
    elif not critical_red and (total_red >= 1 or total_green <= 2):
        key = "2_ready_internal_alignment"
    elif not critical_red and not secondary_red and total_red == 0 and total_green <= 3:
        key = "3_ready_vendor_discovery"
    elif not critical_red and not secondary_red and total_red == 0 and critical_evidence_ok and total_green >= 6:
        key = "7_ready_for_rfp"
    elif not critical_red and total_green >= 5:
        key = "6_ready_enterprise_planning"
    elif not critical_red and not secondary_red and total_green >= 3:
        key = "5_ready_service_line_scale"
    else:
        key = "4_ready_targeted_pilot"
    status = RUBRIC["decision_readiness_statuses"][key]
    warnings = []
    if critical_red:
        warnings.append(RUBRIC["gating_model"]["critical_gating_domains"]["red_override_message"])
    if secondary_red:
        warnings.append(RUBRIC["gating_model"]["secondary_gating_domains"]["red_override_message"])
    return {
        "label": status["label"],
        "guidance": status["guidance"],
        "allowed_next_steps": status["allowed_next_steps"],
        "gating_warnings": warnings,
    }


def score_responses(responses, domain_evidence, phi_confirmed, evidence_analysis=None):
    domain_scores = {}
    for domain_id in DOMAIN_IDS:
        signals = responses.get(domain_id, {})
        has_file = (domain_evidence.get(domain_id) is not None) and phi_confirmed
        if not signals:
            domain_scores[domain_id] = {"readiness_score": 1.0, "evidence_confidence_score": 0.0, "color": "Red"}
            continue
        readiness_scores = [STATED_LEVEL_SCORES.get(str(level), 1.0) for level in signals.values()]
        readiness_avg = sum(readiness_scores) / len(readiness_scores)

        if evidence_analysis and domain_id in evidence_analysis and not evidence_analysis[domain_id].get("error"):
            sig_analyses = evidence_analysis[domain_id].get("signals", {})
            ev_levels = [float(sig_analyses[s]["evidence_level"]) for s in signals if s in sig_analyses]
            evidence_avg = sum(ev_levels) / len(ev_levels) if ev_levels else (3.0 if has_file else 0.0)
        else:
            evidence_avg = 3.0 if has_file else 0.0

        gap = readiness_avg - evidence_avg
        color = _color(readiness_avg)
        if color == "Green" and gap > 2.0:
            color = "Yellow"
        domain_scores[domain_id] = {
            "readiness_score": round(readiness_avg, 2),
            "evidence_confidence_score": round(evidence_avg, 2),
            "color": color,
        }
    return domain_scores, _decision_readiness_status(domain_scores)


def get_anthropic_client():
    try:
        import anthropic
        key = st.secrets.get("ANTHROPIC_API_KEY") or st.secrets.get("anthropic_api_key")
        if key:
            return anthropic.Anthropic(api_key=key)
    except Exception:
        pass
    return None


def extract_file_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    content = uploaded_file.read()
    name = uploaded_file.name.lower()
    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)[:12000]
        elif name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:12000]
        else:
            return content.decode("utf-8", errors="replace")[:12000]
    except Exception as e:
        return f"[Could not read file: {e}]"


def analyze_evidence(client, domain_id, domain_name, gate_type, questions, domain_responses, document_text, filename):
    signal_lines = []
    for q in questions:
        sig_id = q["signal_id"]
        level = domain_responses.get(sig_id, "none")
        opts = SIGNAL_OPTIONS.get(sig_id, [])
        idx = LEVEL_KEYS.index(level) if level in LEVEL_KEYS else 0
        answer = opts[idx] if idx < len(opts) else level
        signal_lines.append(f"- {q['signal_name']} ({sig_id}): \"{answer}\"")

    prompt = f"""You are reviewing a document uploaded to validate self-reported answers in a health system AI readiness assessment.

DOMAIN: {domain_name} ({gate_type})

SELF-REPORTED ANSWERS:
{chr(10).join(signal_lines)}

UPLOADED DOCUMENT (filename: {filename}):
---
{document_text}
---

For each signal listed above, analyze the document and return ONLY a JSON object with this exact structure:

{{
  "domain_id": "{domain_id}",
  "document_filename": "{filename}",
  "document_relevance": "high|medium|low",
  "overall_evidence_quality": <0.0-5.0>,
  "additional_context": "<AI-readiness context from the document not captured in the questionnaire, max 120 words>",
  "signals": {{
    "<signal_id>": {{
      "evidence_level": <0.0-5.0>,
      "corroborates": <true|false>,
      "contradicts": <true|false>,
      "key_quote": "<most relevant quote, max 80 words, or null>",
      "evidence_note": "<one sentence: what the document shows for this signal>"
    }}
  }}
}}

EVIDENCE LEVEL SCALE: 0=no relevant content, 1=tangential, 2=related but weak, 3=relevant and supportive, 4=strong direct evidence, 5=explicit confirmation with metrics or named owners.

Do not make compliance certification claims. Flag contradictions honestly. Return ONLY the JSON object."""

    try:
        import anthropic
        message = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        text = message.content[0].text.strip()
        match = re.search(r'\{.*\}', text, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(text)
    except Exception as e:
        return {"error": str(e), "domain_id": domain_id}


def generate_gamma_prompt(org_name, operator_role, ehr, cloud, domain_scores, decision, responses, context_responses=None, evidence_analysis=None):
    # Target: ≤13 slides (14 with gating) — optimized for Gamma free-plan credit efficiency
    today = date.today().strftime("%B %Y")
    CE = {"Red": "🔴", "Yellow": "🟡", "Green": "🟢"}
    vendors = _select_vendors(ehr, cloud)

    def gate_label(d):
        if d in CRITICAL_GATING: return "Critical"
        if d in SECONDARY_GATING: return "Secondary"
        return "—"

    def ev_label(d):
        return "✅" if domain_scores[d].get("evidence_confidence_score", 0) >= 3 else "⚠️"

    def answer_short(signal_id, level_key):
        opts = SIGNAL_OPTIONS.get(signal_id, [])
        idx = LEVEL_KEYS.index(level_key) if level_key in LEVEL_KEYS else 0
        text = opts[idx] if idx < len(opts) else level_key
        return text[:55] + "…" if len(text) > 55 else text

    def sig_emoji(level):
        return "🟢" if level in ("defined", "specific_with_metrics") else ("🟡" if level == "partial" else "🔴")

    def trunc(text, n):
        return (text[:n] + "…") if text and len(text) > n else (text or "")

    L = []
    def line(s=""): L.append(s)
    def sep(): L.extend(["", "---", ""])

    # ── Theme ──
    line("Create a professional presentation. Dark modern theme. Healthcare enterprise aesthetic. Color-coded indicators. No clip art. Keep each slide concise.")
    sep()

    # ── Slide 1: Title ──
    line(f"# {org_name}")
    line("## AI Infrastructure Readiness Assessment")
    line(f"### {today}" + (f" · {operator_role}" if operator_role else ""))
    sep()

    # ── Slide 2: Overview ──
    label = decision["label"]
    sorted_d = sorted(domain_scores.items(), key=lambda x: x[1]["readiness_score"])
    worst = DOMAINS[sorted_d[0][0]]["name"]
    best  = DOMAINS[sorted_d[-1][0]]["name"]
    line("# Assessment Overview")
    line("")
    line(f"| | |")
    line("|---|---|")
    line(f"| **Status** | **{label}** |")
    line(f"| **EHR** | {ehr} |")
    line(f"| **Cloud** | {cloud} |")
    line(f"| **Strongest domain** | {best} |")
    line(f"| **Weakest domain** | {worst} |")
    line("")
    line(f"> {trunc(decision['guidance'], 160)}")
    sep()

    # ── Slide 3: Scorecard + Credibility Gap (merged to save one slide) ──
    line("# Domain Scorecard & Evidence Confidence")
    line("")
    line("| Domain | Score | Status | Evidence | Gap | Gate |")
    line("|---|---|---|---|---|---|")
    for d in DOMAIN_IDS:
        ds = domain_scores[d]
        gap = ds["readiness_score"] - ds["evidence_confidence_score"]
        if gap > 2.0: interp = "High risk"
        elif gap > 1.2: interp = "Gap"
        elif gap > 0.5: interp = "Validate"
        else: interp = "Supported"
        line(f"| {DOMAINS[d]['name']} | {ds['readiness_score']:.1f} | {CE[ds['color']]} {ds['color']} | {ev_label(d)} {ds['evidence_confidence_score']:.1f} | {interp} | {gate_label(d)} |")
    sep()

    # ── Conditional Slide A: Gating (only if critical Red) ──
    red_crit = [d for d in CRITICAL_GATING if domain_scores.get(d, {}).get("color") == "Red"]
    red_sec  = [d for d in SECONDARY_GATING if domain_scores.get(d, {}).get("color") == "Red"]
    if red_crit:
        line("# ⛔ Gating Warnings — Action Required")
        line("")
        line("No vendor engagement, RFP, or procurement until resolved.")
        line("")
        for d in red_crit:
            line(f"- 🔴 **{DOMAINS[d]['name']}** — Critical Gate · {domain_scores[d]['readiness_score']:.1f}/5.0")
        for d in red_sec:
            line(f"- 🟡 **{DOMAINS[d]['name']}** — Secondary Gate · {domain_scores[d]['readiness_score']:.1f}/5.0")
        for w in decision.get("gating_warnings", []):
            line(f"> ⚠️ {trunc(w, 120)}")
        sep()

    # ── Slides 4–11: One strategic slide per domain ──
    for i, domain_id in enumerate(DOMAIN_IDS):
        ds  = domain_scores[domain_id]
        dom = DOMAINS[domain_id]
        ea  = (evidence_analysis or {}).get(domain_id, {})
        ea_sigs = ea.get("signals", {})
        has_contradiction = any(ea_sigs.get(q["signal_id"], {}).get("contradicts") for q in dom["questions"])
        domain_responses = responses.get(domain_id, {})
        owner = DOMAIN_OWNERS.get(domain_id, "CIO")

        # Classify signals as gaps (none/vague/partial) vs strengths (defined/specific_with_metrics)
        weak_levels = {"none", "vague", "partial"}
        strong_levels = {"defined", "specific_with_metrics"}
        gaps    = [q for q in dom["questions"] if domain_responses.get(q["signal_id"], "none") in weak_levels]
        strengths = [q for q in dom["questions"] if domain_responses.get(q["signal_id"], "none") in strong_levels]

        line(f"# D{i+1}: {dom['name']}")
        line(f"## {gate_label(domain_id)} · {CE[ds['color']]} {ds['color']} {ds['readiness_score']:.1f}/5 · {ev_label(domain_id)} Evidence")
        line("")

        headline = _domain_headline(ds["color"], ds["readiness_score"], domain_id, len(gaps), len(dom["questions"]))
        line(f"**{headline}**")
        line("")

        if ds["color"] in ("Red", "Yellow") and gaps:
            # Gaps with signal-level evidence context
            line("**Gaps requiring action:**")
            for q in gaps[:3]:
                sid   = q["signal_id"]
                level = domain_responses.get(sid, "none")
                ans   = answer_short(sid, level)
                sig_e = ea_sigs.get(sid, {})
                ev_flag = ""
                if sig_e.get("contradicts"):
                    ev_flag = " ⚠️ *contradicts uploaded evidence*"
                elif sig_e.get("evidence_note"):
                    ev_flag = f" · *{trunc(sig_e['evidence_note'], 55)}*"
                line(f"- {sig_emoji(level)} **{q['signal_name']}** — {ans}{ev_flag}")
            line("")

            # Reference design — gap closure
            for v in vendors:
                cap = REF_DOMAIN_CAPABILITIES.get(v, {}).get(domain_id, "")
                if cap:
                    line(f"**{REF_VENDOR_NAMES[v]} — industry reference:**")
                    line(f"> {cap}")
                    line("")

            # Owner and action
            line(f"**Owner:** {owner}")

        else:
            # Green domain — compact strengths summary
            if strengths:
                names = " · ".join(q["signal_name"] for q in strengths[:4])
                line(f"**Capabilities confirmed:** {names}")
            line("")
            line("No immediate action required. Maintain current practices and validate against AI workload scaling requirements as use cases mature.")

        # Evidence additional context (AI document analysis)
        if ea.get("additional_context"):
            line("")
            line(f"> 🔍 *Evidence note: {trunc(ea['additional_context'], 110)}*")

        # Context question (reviewer note, not scored)
        if context_responses:
            for cid, cq in CONTEXT_QUESTIONS.items():
                if cq["domain"] == domain_id and cid in context_responses:
                    line(f"> 📋 *{cq['stem'].replace('For context only (not scored): ', '')} — {context_responses[cid]}*")

        if has_contradiction:
            line("> ⚠️ *Contradiction detected between self-report and uploaded evidence — verify before procurement use.*")

        sep()

    # ── Slide 12: Reference Design Match — Strategic Insights to Action ──
    weak_domains = [d for d in DOMAIN_IDS if domain_scores.get(d, {}).get("color") in ("Red", "Yellow")]
    red_domains  = [d for d in DOMAIN_IDS if domain_scores.get(d, {}).get("color") == "Red"]
    green_domains = [d for d in DOMAIN_IDS if domain_scores.get(d, {}).get("color") == "Green"]
    vnames = " · ".join(REF_VENDOR_NAMES[v] for v in vendors)

    line("# From Assessment to Action — Reference Design Findings")
    line(f"## {vnames} · Matched to {ehr} · {cloud}")
    line("")

    # Vendor fit rationale (strategic, not a table)
    vendor_strategic = {
        "epic":   f"As an active Epic customer, the path of least resistance is activating capabilities already embedded in your Epic environment — App Orchard AI, native cognitive models, and DAX Copilot — before engaging external vendors.",
        "aws":    f"With AWS as your cloud platform, HealthLake can serve as the FHIR-native data layer that feeds AI workloads, and Bedrock provides the generative AI infrastructure for ambient documentation and summarization use cases.",
        "azure":  f"Your Microsoft alignment makes DAX Copilot the highest-readiness ambient documentation path and Azure Health Data Services the natural FHIR aggregation layer — both may be negotiable within your existing EA.",
        "google": f"Google Cloud's Healthcare Data Engine is the strongest available solution for FHIR-native data harmonization and de-identification — priority activation if data readiness is the primary barrier to AI workload deployment.",
    }
    for v in vendors:
        line(f"**{REF_VENDOR_NAMES[v]}:** {vendor_strategic[v]}")
        line("")

    if weak_domains:
        n_red = len(red_domains)
        n_weak = len(weak_domains)
        n_green = len(green_domains)
        line(f"**Assessment verdict:** {n_red} domain(s) are blocking · {n_weak - n_red} require targeted maturation · {n_green} are ready to proceed")
        line("")
        line("**Priority action map — close these gaps to unlock AI procurement:**")
        line("")
        line("| Priority | Domain | Gap | Industry Reference Action | Owner |")
        line("|---|---|---|---|---|")
        def _gap_cell(d, base_label, score):
            ea_d = (evidence_analysis or {}).get(d, {})
            has_contradiction = any(s.get("contradicts") for s in ea_d.get("signals", {}).values())
            if has_contradiction:
                return f"⚠️ Evidence contradiction — {score:.1f}/5"
            ctx = ea_d.get("additional_context", "")
            if ctx:
                return trunc(ctx, 65)
            return f"{base_label} — {score:.1f}/5"

        pri = 1
        for d in red_domains:
            ds_d = domain_scores[d]
            v = vendors[0]
            cap = REF_DOMAIN_CAPABILITIES.get(v, {}).get(d, "")
            act = trunc(cap, 90) if cap else "Address foundational gap before procurement"
            line(f"| 🔴 P{pri} | {DOMAINS[d]['name']} | {_gap_cell(d, 'Blocking', ds_d['readiness_score'])} | {act} | {DOMAIN_OWNERS.get(d, 'CIO')} |")
            pri += 1
        for d in weak_domains:
            if d in red_domains:
                continue
            ds_d = domain_scores[d]
            v = vendors[0]
            cap = REF_DOMAIN_CAPABILITIES.get(v, {}).get(d, "")
            act = trunc(cap, 90) if cap else "Close gap before vendor engagement"
            line(f"| 🟡 P{pri} | {DOMAINS[d]['name']} | {_gap_cell(d, 'Needs maturation', ds_d['readiness_score'])} | {act} | {DOMAIN_OWNERS.get(d, 'CIO')} |")
            pri += 1
        if len(vendors) > 1:
            line("")
            line(f"*Secondary reference: {REF_VENDOR_NAMES[vendors[1]]} — solutions noted on each domain slide above.*")
    else:
        line("**All domains are Green.** This organization is prepared to proceed with AI vendor engagement and procurement.")
        line("")
        line("Recommended next step: issue RFP to matched reference design vendors and begin formal vendor evaluation.")

    line("")
    line(f"*BT Compass Reference Library v1.0 · {', '.join(REF_VENDOR_FILES[v] for v in vendors)} · For internal planning only. Not procurement advice.*")
    sep()

    # ── Slide 13: Recommended Path + Disclaimer (merged to save one slide) ──
    line("# Recommended Path")
    line(f"## {label}")
    line("")
    line(trunc(decision["guidance"], 200))
    line("")
    line("**Allowed next steps:**")
    for step in decision.get("allowed_next_steps", []):
        line(f"- {step}")
    line("")
    line("---")
    line(f"*{org_name} · BT Compass · {today} · For internal planning only. Not legal, clinical, or procurement advice.*")

    return "\n".join(L)


# ── RFP Word Document Generation ────────────────────────────────────────────

def _extract_ref_section(md_text, section_name):
    """Extract content of a ## section from a markdown file, up to the next ## or end."""
    pattern = rf'## {re.escape(section_name)}\n(.*?)(?=\n## |\Z)'
    match = re.search(pattern, md_text, re.DOTALL)
    return match.group(1).strip() if match else ""


def _md_to_docx(doc, md_text):
    """Parse markdown (### headings, - bullets, plain paragraphs) into docx content."""
    for raw_line in md_text.split('\n'):
        line = raw_line.strip()
        if not line or line == '---':
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)


def generate_rfp_docx(org_name, operator_role, ehr, cloud, domain_scores, decision, vendors, evidence_analysis=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    today_str = date.today().strftime("%B %d, %Y")

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def add_colored_heading(text, level, hex_color="1F3864"):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(hex_color)
        return h

    def add_divider():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_label_value(label, value):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    # ── Cover ──
    title = doc.add_heading(org_name, level=0)
    title.runs[0].font.color.rgb = RGBColor.from_string("1F3864")

    sub = doc.add_heading("AI Infrastructure Readiness — Procurement Requirements Brief", level=1)
    sub.runs[0].font.color.rgb = RGBColor.from_string("2E74B5")

    p = doc.add_paragraph()
    p.add_run(f"BT Compass Assessment  ·  {today_str}")
    if operator_role:
        p.add_run(f"  ·  {operator_role}")

    doc.add_paragraph()
    add_label_value("Decision-Readiness Status", decision["label"])
    add_label_value("EHR Platform", ehr)
    add_label_value("Cloud Environment", cloud)
    add_label_value("Reference Designs", "  ·  ".join(REF_VENDOR_NAMES[v] for v in vendors))
    doc.add_paragraph()

    p = doc.add_paragraph(decision["guidance"])
    p.runs[0].italic = True

    doc.add_paragraph()
    disc = doc.add_paragraph(
        "For internal IT and procurement use only. Not legal, clinical, compliance, or final procurement advice. "
        "All requirements should be reviewed with legal and compliance before external use."
    )
    disc.runs[0].font.size = Pt(9)
    disc.runs[0].font.color.rgb = RGBColor.from_string("666666")

    doc.add_page_break()

    # ── Section 1: Procurement Priority Summary ──
    add_colored_heading("Section 1 — Procurement Priority Summary", level=1)
    doc.add_paragraph(
        "The following domains have been identified as requiring action before AI vendor engagement or procurement can proceed. "
        "Domains are listed in priority order — Red (blocking) first, then Yellow (requiring maturation)."
    )
    doc.add_paragraph()

    red_domains  = [d for d in DOMAIN_IDS if domain_scores[d]["color"] == "Red"]
    yellow_domains = [d for d in DOMAIN_IDS if domain_scores[d]["color"] == "Yellow"]
    green_domains  = [d for d in DOMAIN_IDS if domain_scores[d]["color"] == "Green"]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Domain", "Score", "Status", "Gate Type"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for d in red_domains + yellow_domains + green_domains:
        ds = domain_scores[d]
        row = table.add_row().cells
        row[0].text = DOMAINS[d]["name"]
        row[1].text = f"{ds['readiness_score']:.1f} / 5.0"
        row[2].text = ds["color"]
        if d in CRITICAL_GATING:
            row[3].text = "Critical Gate"
        elif d in SECONDARY_GATING:
            row[3].text = "Secondary Gate"
        else:
            row[3].text = "—"

    doc.add_paragraph()
    if red_domains:
        p = doc.add_paragraph()
        p.add_run("⛔ Blocking domains: ").bold = True
        p.add_run(
            f"{', '.join(DOMAINS[d]['name'] for d in red_domains)} — "
            "no vendor engagement, RFP, or procurement should proceed until these are resolved."
        )

    # Evidence documents reviewed (if any uploads were analyzed)
    analyzed = {d: ea for d, ea in (evidence_analysis or {}).items() if not ea.get("error")}
    if analyzed:
        doc.add_paragraph()
        add_colored_heading("Evidence Documents Reviewed", level=2, hex_color="2E74B5")
        doc.add_paragraph(
            "The following documents were uploaded and analyzed by Claude AI during the assessment. "
            "Evidence findings are incorporated into the domain-specific requirements in Section 3."
        )
        ev_table = doc.add_table(rows=1, cols=4)
        ev_table.style = 'Table Grid'
        ev_hdr = ev_table.rows[0].cells
        for i, h in enumerate(["Domain", "Document", "Quality", "Relevance"]):
            ev_hdr[i].text = h
            ev_hdr[i].paragraphs[0].runs[0].bold = True
        for d, ea in analyzed.items():
            ev_row = ev_table.add_row().cells
            ev_row[0].text = DOMAINS[d]["name"]
            ev_row[1].text = ea.get("document_filename", "uploaded file")
            ev_row[2].text = f"{ea.get('overall_evidence_quality', 0):.1f} / 5.0"
            ev_row[3].text = ea.get("document_relevance", "—").capitalize()
        contradicting = [d for d, ea in analyzed.items()
                         if any(s.get("contradicts") for s in ea.get("signals", {}).values())]
        if contradicting:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("⚠️ Contradictions detected: ").bold = True
            p.add_run(
                f"{', '.join(DOMAINS[d]['name'] for d in contradicting)} — "
                "uploaded documents contain information that conflicts with self-reported scores. "
                "Review the specific findings in Section 3 before using these scores for procurement decisions."
            )
    add_divider()

    # ── Section 2: Vendor Reference Requirements ──
    doc.add_paragraph()
    add_colored_heading("Section 2 — Vendor Reference Requirements", level=1)
    doc.add_paragraph(
        "The following requirements are derived from the BT Compass Reference Design Library. "
        "They reflect standard procurement practice for health systems engaging these vendor ecosystems. "
        "Requirements should be tailored to the specific products and scope of your RFP."
    )

    for v in vendors:
        doc.add_paragraph()
        add_colored_heading(REF_VENDOR_NAMES[v], level=2, hex_color="2E74B5")
        ref_text = (BASE / "reference_library" / REF_VENDOR_FILES[v]).read_text()
        proc_section = _extract_ref_section(ref_text, "typical_procurement_requirements")
        if proc_section:
            _md_to_docx(doc, proc_section)
        else:
            doc.add_paragraph("Procurement requirements not available for this vendor.")

    add_divider()

    # ── Section 3: Domain-by-Domain RFP Requirements ──
    doc.add_paragraph()
    add_colored_heading("Section 3 — Domain-by-Domain RFP Requirements", level=1)
    doc.add_paragraph(
        "The following requirements correspond to domains where assessment scores indicate gaps. "
        "Include these as requirements in vendor RFPs, SOWs, or due diligence questionnaires for the specific domain. "
        "Green domains are omitted — no additional requirements beyond standard vendor qualification apply."
    )

    priority_domains = red_domains + yellow_domains
    if not priority_domains:
        doc.add_paragraph("All domains scored Green. No domain-specific gap requirements to include.")
    else:
        for d in priority_domains:
            ds = domain_scores[d]
            rfp = RFP_DOMAIN_REQUIREMENTS.get(d)
            if not rfp:
                continue
            doc.add_paragraph()
            gate = "Critical Gate" if d in CRITICAL_GATING else ("Secondary Gate" if d in SECONDARY_GATING else "Standard")
            heading_text = f"{DOMAINS[d]['name']}  ·  {ds['color']} {ds['readiness_score']:.1f}/5  ·  {gate}"
            add_colored_heading(heading_text, level=2, hex_color="2E74B5")

            # Evidence findings for this domain (if a document was uploaded and analyzed)
            ea_d = (evidence_analysis or {}).get(d, {})
            if ea_d and not ea_d.get("error"):
                filename = ea_d.get("document_filename", "uploaded document")
                add_colored_heading(f"Evidence Findings — {filename}", level=3, hex_color="1F5C1F")
                if ea_d.get("additional_context"):
                    doc.add_paragraph(ea_d["additional_context"])
                # Signal-level evidence notes
                ea_sigs = ea_d.get("signals", {})
                contradictions = [(sid, s) for sid, s in ea_sigs.items() if s.get("contradicts")]
                corroborations = [(sid, s) for sid, s in ea_sigs.items() if s.get("corroborates") and s.get("key_quote")]
                if contradictions:
                    p = doc.add_paragraph()
                    p.add_run("⚠️ Contradictions detected — procurement risk:").bold = True
                    for sid, s in contradictions[:3]:
                        note = s.get("evidence_note") or s.get("key_quote") or ""
                        doc.add_paragraph(f"{sid}: {note}", style='List Bullet')
                if corroborations:
                    p = doc.add_paragraph()
                    p.add_run("Evidence corroborates:").bold = True
                    for sid, s in corroborations[:3]:
                        quote = s.get("key_quote", "")
                        if quote:
                            doc.add_paragraph(f'"{quote[:120]}"', style='List Bullet')
                doc.add_paragraph()

            add_colored_heading(rfp["heading"], level=3, hex_color="404040")
            for item in rfp["items"]:
                doc.add_paragraph(item, style='List Bullet')

    add_divider()

    # ── Section 4: Standard Health AI Contract Requirements ──
    doc.add_paragraph()
    add_colored_heading("Section 4 — Standard Health AI Contract Requirements", level=1)
    doc.add_paragraph(
        "The following requirements apply to all health AI vendor engagements regardless of domain or vendor. "
        "These should be included as baseline contract terms in any AI vendor agreement."
    )

    std_sections = {
        "HIPAA Business Associate Agreement (BAA)": [
            "BAA must be executed before any PHI is shared, including during proof-of-concept.",
            "BAA must identify all permitted uses of PHI and explicitly prohibit use of health system PHI for AI model training without explicit written consent.",
            "BAA must include a complete, current subprocessor list as an exhibit, with obligation to notify health system of material subprocessor changes.",
            "BAA must specify data residency requirements and confirm US-only storage and processing if required.",
        ],
        "Data Security & Breach Notification": [
            "Vendor must notify the health system of any confirmed or suspected breach within 72 hours of discovery.",
            "Breach notification must include: nature of the incident, PHI categories and approximate record count affected, steps taken to mitigate, and corrective action plan.",
            "Vendor must maintain cyber liability insurance of not less than $5M per occurrence (confirm with your legal and risk team for appropriate threshold).",
            "Vendor must provide the health system with the right to conduct or commission a security assessment annually.",
        ],
        "Service Level Agreement (SLA) Minimums": [
            "Production uptime: minimum 99.5% monthly availability for AI services in active clinical or operational use.",
            "Incident response: P1 (system down) — 1-hour acknowledgment, 4-hour resolution target.",
            "Planned maintenance windows: minimum 5 business days advance notice; no maintenance during business hours without health system approval.",
            "SLA credits: vendor must provide service credits for failures to meet SLA targets — credits must be contractually enforceable.",
        ],
        "Audit Rights & Compliance": [
            "Health system retains the right to audit vendor's data handling and security practices on reasonable notice (not less than 30 days).",
            "Vendor must provide evidence of annual third-party security audit or penetration test upon request.",
            "Vendor must maintain audit logs of all PHI access for a minimum of 6 years.",
            "Vendor must cooperate with any regulatory audit or investigation related to the health system's use of the vendor's AI services.",
        ],
        "Contract Termination & Data Portability": [
            "Upon contract termination, vendor must return or destroy all PHI within 30 days and provide written certification of destruction.",
            "Health system retains ownership of all data provided to the vendor and all outputs generated from that data.",
            "Vendor must provide data export in a standard, machine-readable format upon request — export must be completed within 30 days.",
            "Termination for cause provisions must include the right to terminate immediately upon confirmed data breach or material BAA violation.",
        ],
    }

    for section_title, items in std_sections.items():
        doc.add_paragraph()
        add_colored_heading(section_title, level=3, hex_color="2E74B5")
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # ── Disclaimer ──
    doc.add_page_break()
    add_colored_heading("Scope & Disclaimer", level=2, hex_color="666666")
    disclaimer_items = [
        "This document is for internal IT and procurement planning purposes only.",
        "It is not legal advice, compliance certification, or final procurement guidance.",
        "All requirements should be reviewed and approved by legal counsel and compliance before external use.",
        "Vendor capabilities described in reference designs reflect publicly available information as of the BT Compass Library update date. Validate against current vendor documentation before procurement.",
        f"BT Compass AI Infrastructure Assessment  ·  {today_str}  ·  breakthroughcompass.com",
    ]
    for item in disclaimer_items:
        doc.add_paragraph(item, style='List Bullet')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── UI ──────────────────────────────────────────────────────────────────────

st.set_page_config(page_title="BT Compass Assessment", page_icon="🧭", layout="wide")

st.title("🧭 BT Compass AI Infrastructure Assessment")
st.caption("Health AI Readiness for CIOs · breakthroughcompass.com")

with st.sidebar:
    st.header("About Your Organization")
    org_name = st.text_input("Health System Name", placeholder="e.g. Atlantic Health System")
    operator_role = st.text_input("Your Role", placeholder="e.g. VP of IT Infrastructure")
    ehr = st.selectbox("Primary EHR", ["Epic", "Cerner / Oracle Health", "Meditech", "CPSI", "Other"])
    cloud = st.selectbox("Primary Cloud", ["Azure", "AWS", "Google Cloud", "On-premise", "Hybrid / Mixed"])
    st.divider()
    st.caption(
        "BT Compass scores AI readiness across 8 domains using a deterministic rubric "
        "aligned to NIST AI RMF, WHO Ethics, and CHAI Blueprint standards."
    )

DOMAINS = _build_domain_questions()
COLOR_EMOJI = {"Red": "🔴", "Yellow": "🟡", "Green": "🟢"}

st.markdown("### Complete each section, then click **Run Assessment** at the bottom.")
st.markdown("*Critical Gate domains (marked 🔴) can block decision-readiness regardless of other scores.*")

# PHI attestation — required before evidence uploads are enabled
st.markdown("---")
phi_confirmed = st.checkbox(
    f"**Evidence upload attestation:** {PHI_ATTESTATION_TEXT}",
    value=False,
    help="You must confirm this before uploading any supporting documents.",
)
if phi_confirmed:
    st.success("Attestation confirmed. You may upload supporting evidence in each domain below.")
else:
    st.info("Check the box above to enable evidence uploads. Evidence improves your score confidence.")
st.markdown("---")

responses = {}
domain_evidence = {}
context_responses = {}

CONTEXT_QUESTIONS = {
    "s4_incident_history": {
        "stem": "For context only (not scored): Has your organization experienced a significant cybersecurity incident in the past 24 months?",
        "options": [
            "No significant incident in the past 24 months.",
            "Yes — a significant incident occurred and has been remediated.",
            "Yes — a significant incident occurred and remediation is ongoing.",
            "Prefer not to disclose.",
        ],
        "domain": "security_privacy_compliance",
    },
}

for domain_id in DOMAIN_IDS:
    dom = DOMAINS[domain_id]
    if domain_id in CRITICAL_GATING:
        badge = "  🔴 *Critical Gate*"
    elif domain_id in SECONDARY_GATING:
        badge = "  🟡 *Secondary Gate*"
    else:
        badge = ""

    with st.expander(f"**{dom['name']}**{badge}", expanded=True):
        responses[domain_id] = {}

        for q in dom["questions"]:
            options = SIGNAL_OPTIONS.get(q["signal_id"], [
                "Not in place",
                "Informal — discussed but not documented",
                "Partially in place — some elements exist",
                "Formally defined and documented",
                "Formally defined with measurable targets or metrics",
            ]) + [NA_OPTION]
            choice_idx = st.selectbox(
                f"**{q['signal_name']}**  \n*{q['hint']}*",
                range(len(options)),
                format_func=lambda i, opts=options: opts[i],
                key=f"{domain_id}_{q['signal_id']}",
            )
            responses[domain_id][q["signal_id"]] = LEVEL_KEYS[choice_idx] if choice_idx < len(LEVEL_KEYS) else "none"

        # Context-only questions for this domain (not scored, not gated)
        for cid, cq in CONTEXT_QUESTIONS.items():
            if cq["domain"] == domain_id:
                st.markdown("---")
                st.caption("ℹ️ Context question — recorded for reviewer use only. Does not affect your score or readiness status.")
                ctx_idx = st.selectbox(
                    f"*{cq['stem']}*",
                    range(len(cq["options"])),
                    format_func=lambda i, opts=cq["options"]: opts[i],
                    key=f"ctx_{cid}",
                )
                context_responses[cid] = cq["options"][ctx_idx]

        # Evidence upload for this domain
        st.markdown("---")
        all_evidence_types = []
        for q in dom["questions"]:
            all_evidence_types.extend(q["evidence_types"])
        unique_types = list(dict.fromkeys(all_evidence_types))

        if phi_confirmed:
            uploaded = st.file_uploader(
                "Upload supporting evidence for this domain (optional — redacted documents only)",
                key=f"upload_{domain_id}",
                accept_multiple_files=False,
                help="Accepted: " + " · ".join(unique_types[:4]),
            )
            domain_evidence[domain_id] = uploaded
            if unique_types:
                st.caption("Accepted evidence: " + "  ·  ".join(unique_types))
        else:
            st.caption("*Enable evidence uploads by confirming the attestation above.*")
            domain_evidence[domain_id] = None

st.markdown("---")

if st.button("🧭 Run Assessment", type="primary", use_container_width=True):
    if not org_name:
        st.warning("Please enter your Health System Name in the sidebar before running.")
    else:
        # Analyze uploaded documents with Claude before scoring
        evidence_analysis = {}
        uploads = {d: f for d, f in domain_evidence.items() if f is not None}
        if uploads and phi_confirmed:
            client = get_anthropic_client()
            if client:
                with st.spinner(f"Analyzing {len(uploads)} uploaded document(s) with Claude…"):
                    for domain_id, uploaded_file in uploads.items():
                        dom = DOMAINS[domain_id]
                        gate = "Critical Gate" if domain_id in CRITICAL_GATING else ("Secondary Gate" if domain_id in SECONDARY_GATING else "Standard")
                        text = extract_file_text(uploaded_file)
                        evidence_analysis[domain_id] = analyze_evidence(
                            client, domain_id, dom["name"], gate,
                            dom["questions"], responses.get(domain_id, {}),
                            text, uploaded_file.name,
                        )
            else:
                st.info("ℹ️ Add your Anthropic API key to Streamlit secrets to enable AI document analysis. Documents registered as uploaded — evidence confidence set to default.")

        domain_scores, decision = score_responses(responses, domain_evidence, phi_confirmed, evidence_analysis)
        st.session_state["results"] = {
            "org_name": org_name, "operator_role": operator_role, "ehr": ehr, "cloud": cloud,
            "domain_scores": domain_scores, "decision": decision, "responses": responses,
            "context_responses": context_responses, "evidence_analysis": evidence_analysis,
        }

        st.markdown(f"## Results — {org_name}")
        if operator_role:
            st.markdown(f"*{operator_role}  ·  {ehr}  ·  {cloud}*")

        # Decision-readiness status
        st.markdown("### Decision-Readiness Status")
        label = decision["label"]
        guidance = decision["guidance"]

        not_ready_labels = {"Not Ready", "Ready for Internal Alignment"}
        if label in not_ready_labels:
            st.error(f"**{label}**\n\n{guidance}")
        elif "Discovery" in label or "Pilot" in label:
            st.warning(f"**{label}**\n\n{guidance}")
        else:
            st.success(f"**{label}**\n\n{guidance}")

        for warning in decision.get("gating_warnings", []):
            st.warning(f"⚠️ {warning}")

        col_steps, col_spacer = st.columns([2, 1])
        with col_steps:
            st.markdown("**Recommended Next Steps:**")
            for step in decision.get("allowed_next_steps", []):
                st.markdown(f"- {step}")

        # Domain scorecard
        st.markdown("---")
        st.markdown("### Domain Scorecard")

        header = st.columns([4, 1, 1, 1, 2])
        header[0].markdown("**Domain**")
        header[1].markdown("**Score**")
        header[2].markdown("**Status**")
        header[3].markdown("**Evidence**")
        header[4].markdown("**Gate Type**")
        st.markdown("---")

        for domain_id in DOMAIN_IDS:
            ds = domain_scores[domain_id]
            dom = DOMAINS[domain_id]
            color = ds["color"]
            emoji = COLOR_EMOJI[color]
            ev_score = ds.get("evidence_confidence_score", 0.0)
            ev_label = "Supported" if ev_score >= 3.0 else "Self-reported"

            if domain_id in CRITICAL_GATING:
                gate = "Critical"
            elif domain_id in SECONDARY_GATING:
                gate = "Secondary"
            else:
                gate = "—"

            row = st.columns([4, 1, 1, 1, 2])
            row[0].write(dom["name"])
            row[1].write(f"{ds['readiness_score']:.1f} / 5.0")
            row[2].write(f"{emoji} {color}")
            row[3].write(ev_label)
            row[4].write(gate)

        st.markdown("---")
        st.caption(
            "Scored by BT Compass deterministic rubric engine · "
            "NIST AI RMF · WHO Ethics · CHAI Blueprint · breakthroughcompass.com"
        )

        # Evidence analysis results — shown per domain when Claude analyzed a document
        if evidence_analysis:
            st.markdown("---")
            st.markdown("#### Evidence Analysis *(Claude-reviewed documents)*")
            for domain_id, ea in evidence_analysis.items():
                if ea.get("error"):
                    st.warning(f"**{DOMAINS[domain_id]['name']}:** Analysis failed — {ea['error']}")
                    continue
                dom_name = DOMAINS[domain_id]["name"]
                relevance = ea.get("document_relevance", "unknown")
                quality = ea.get("overall_evidence_quality", 0)
                filename = ea.get("document_filename", "uploaded file")
                contradictions = [sid for sid, s in ea.get("signals", {}).items() if s.get("contradicts")]
                if contradictions:
                    st.error(f"**{dom_name}** — `{filename}` · Quality: {quality:.1f}/5.0 · Relevance: {relevance} · ⚠️ Contradictions detected")
                else:
                    st.success(f"**{dom_name}** — `{filename}` · Quality: {quality:.1f}/5.0 · Relevance: {relevance}")
                if ea.get("additional_context"):
                    st.caption(f"📌 {ea['additional_context']}")

        # Reviewer context — not scored, not gated
        if context_responses:
            st.markdown("---")
            st.markdown("#### Reviewer Context *(not scored)*")
            for cid, answer in context_responses.items():
                cq = CONTEXT_QUESTIONS[cid]
                st.info(f"**{cq['stem']}**\n\n{answer}")

# ── Gamma Button — shown after any completed assessment ──────────────────────
if "results" in st.session_state:
    r = st.session_state["results"]
    st.markdown("---")
    st.markdown("### Generate Board Presentation")
    st.markdown("Paste the output below directly into **[Gamma.app](https://gamma.app)** to generate a slide deck from your assessment results.")
    if st.button("📊 Generate Gamma Presentation Outline", use_container_width=True):
        gamma_text = generate_gamma_prompt(
            r["org_name"], r["operator_role"], r["ehr"], r["cloud"],
            r["domain_scores"], r["decision"], r["responses"],
            r.get("context_responses", {}), r.get("evidence_analysis", {}),
        )
        st.text_area("Gamma-ready outline — copy and paste into Gamma.app", gamma_text, height=400)
        st.caption("In Gamma.app: click 'Create new' → 'Paste in text' → paste this outline → Generate.")

    # ── RFP Requirements Brief ──
    st.markdown("---")
    st.markdown("### Download RFP Requirements Brief")
    st.markdown(
        "Word document for IT and Procurement teams — vendor-specific requirements derived from the "
        "BT Compass Reference Design Library, domain-by-domain RFP language for each gap area, "
        "and standard health AI contract requirements."
    )
    try:
        rfp_vendors = _select_vendors(r["ehr"], r["cloud"])
        rfp_bytes = generate_rfp_docx(
            r["org_name"], r["operator_role"], r["ehr"], r["cloud"],
            r["domain_scores"], r["decision"], rfp_vendors,
            evidence_analysis=r.get("evidence_analysis", {}),
        )
        org_slug = "".join(c if c.isalnum() else "_" for c in r["org_name"])[:30]
        st.download_button(
            label="⬇️ Download RFP Requirements Brief (.docx)",
            data=rfp_bytes,
            file_name=f"{org_slug}_RFP_Requirements.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.caption(
            "Includes: vendor procurement requirements · domain-by-domain RFP language · "
            "standard health AI contract terms (BAA, SLA, audit rights, data portability)."
        )
    except Exception as e:
        st.error(f"Could not generate RFP document: {e}")
